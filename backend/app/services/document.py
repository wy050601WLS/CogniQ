"""文件服务"""
import os
import uuid
import asyncio
import logging
from fastapi import UploadFile

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.config import settings
from app.models.document import Document, DocumentVersion
from app.models.user_file import UserFile
from app.models.tag import Tag
from app.utils.helpers import get_or_404
from app.exceptions import ValidationError, ProcessingError, PermissionDeniedError, NotFoundError

logger = logging.getLogger(__name__)

# 支持的文件类型
SUPPORTED_FILE_TYPES = {"pdf", "docx", "doc", "md", "txt", "html", "htm"}


class DocumentService:
    """文件业务逻辑"""

    @staticmethod
    async def get_by_id(db: AsyncSession, doc_id: str, user_id: str = None) -> Document:
        """获取文件详情"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 验证访问权限：所有者、公开文件、或有引用关系
        if user_id:
            has_access = (
                doc.owner_id == user_id or
                doc.is_public or
                await DocumentService._has_user_file(db, user_id, doc_id)
            )
            if not has_access:
                raise PermissionDeniedError("无权访问此文件")

        return doc

    @staticmethod
    async def _has_user_file(db: AsyncSession, user_id: str, doc_id: str) -> bool:
        """检查用户是否有文件引用"""
        result = await db.execute(
            select(UserFile).where(
                UserFile.user_id == user_id,
                UserFile.document_id == doc_id
            )
        )
        return result.scalar_one_or_none() is not None

    @staticmethod
    async def _get_user_file(db: AsyncSession, user_id: str, doc_id: str) -> UserFile:
        """获取用户文件引用"""
        result = await db.execute(
            select(UserFile).where(
                UserFile.user_id == user_id,
                UserFile.document_id == doc_id
            )
        )
        return result.scalar_one_or_none()

    @staticmethod
    async def upload_and_process(
        db: AsyncSession, file: UploadFile, user_id: str, description: str = None, tag_ids: list[str] = None
    ) -> Document:
        """上传并处理文件"""
        # 验证文件
        if not file.filename:
            raise ValidationError("文件名不能为空")

        # 验证文件类型
        file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        if file_ext not in SUPPORTED_FILE_TYPES:
            raise ValidationError(
                f"不支持的文件类型: {file_ext}。支持的类型: {', '.join(sorted(SUPPORTED_FILE_TYPES))}"
            )

        # 检查文件大小
        content = await file.read()
        max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
        if len(content) > max_size:
            raise ValidationError(f"文件大小超过限制（最大 {settings.MAX_UPLOAD_SIZE_MB}MB）")

        # 保存文件
        file_name = f"{uuid.uuid4()}.{file_ext}"
        file_path = os.path.join(settings.UPLOAD_DIR, file_name)

        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        try:
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logger.error(f"文件保存失败: {e}")
            raise ProcessingError("文件保存失败")

        # 创建文件记录
        doc = Document(
            owner_id=user_id,
            filename=file.filename,
            description=description,
            file_type=file_ext,
            file_size=len(content),
            file_path=file_path,
            status="processing",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        # 处理标签
        if tag_ids:
            from app.models.tag import Tag, document_tag_table
            for tag_id in tag_ids:
                tag = await db.get(Tag, tag_id)
                if tag:
                    await db.execute(
                        document_tag_table.insert().values(
                            document_id=doc.id,
                            tag_id=tag_id
                        )
                    )
            await db.commit()

        # 处理文件（线程池避免阻塞事件循环）
        try:
            from app.services.chat import process_document

            chunk_count = await asyncio.to_thread(
                process_document,
                kb_id=doc.id,
                doc_id=doc.id,
                file_path=file_path,
                filename=file.filename,
            )
            doc.status = "completed"
            doc.chunk_count = chunk_count
            logger.info(f"文件处理完成: {file.filename}, 分块数: {chunk_count}")
        except Exception as e:
            doc.status = "error"
            error_msg = str(e)
            doc.error_message = error_msg[:500] + "..." if len(error_msg) > 500 else error_msg
            logger.error(f"文件处理失败: {e}")

        await db.commit()
        await db.refresh(doc)
        return doc

    @staticmethod
    async def update_info(
        db: AsyncSession, doc_id: str, data, user_id: str
    ) -> Document:
        """更新文件信息"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 验证所有权
        if doc.owner_id != user_id:
            raise PermissionDeniedError("无权修改此文件")

        # 验证是否为复制文件（复制文件无修改权）
        if doc.is_copied:
            raise PermissionDeniedError("此文件为复制文件，无法修改")

        # 更新字段
        if data.filename is not None:
            doc.filename = data.filename
        if data.description is not None:
            doc.description = data.description
        if data.is_public is not None:
            doc.is_public = data.is_public

        # 更新标签
        if data.tag_ids is not None:
            # 清除旧标签
            doc.tags.clear()
            # 添加新标签
            for tag_id in data.tag_ids:
                tag = await db.get(Tag, tag_id)
                if tag:
                    doc.tags.append(tag)

        await db.commit()
        await db.refresh(doc)
        return doc

    @staticmethod
    async def delete(db: AsyncSession, doc_id: str, user_id: str) -> None:
        """删除文件"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 检查是否为文件所有者
        is_owner = doc.owner_id == user_id

        # 检查是否有引用关系
        user_file = await DocumentService._get_user_file(db, user_id, doc_id)

        # 如果不是所有者且没有引用关系，则无权删除
        if not is_owner and not user_file:
            raise PermissionDeniedError("无权删除此文件")

        # 如果是引用文件（非所有者），只移除引用关系
        if user_file and not is_owner:
            await db.delete(user_file)
            doc.copy_count = max(0, doc.copy_count - 1)
            await db.commit()
            logger.info(f"文件引用已移除: {doc_id} <- 用户 {user_id}")
            return

        # 删除向量数据
        try:
            from app.services.vector_store import get_vector_store
            vector_store = get_vector_store()
            vector_store.delete_documents(doc_id)
        except Exception as e:
            logger.warning(f"删除向量数据失败: {e}")

        # 删除版本历史
        versions = await db.execute(
            select(DocumentVersion).where(DocumentVersion.document_id == doc_id)
        )
        for version in versions.scalars().all():
            if version.file_path and os.path.exists(version.file_path):
                try:
                    os.remove(version.file_path)
                except Exception:
                    pass
            await db.delete(version)

        # 删除磁盘上的文件
        if doc.file_path:
            try:
                if os.path.exists(doc.file_path):
                    os.remove(doc.file_path)
                    logger.info(f"已删除磁盘文件: {doc.file_path}")
            except Exception as e:
                logger.warning(f"删除磁盘文件失败: {e}")

        await db.delete(doc)
        await db.commit()

    @staticmethod
    async def replace_file(
        db: AsyncSession, doc_id: str, file: UploadFile, user_id: str
    ) -> Document:
        """替换文件内容（创建新版本）"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 验证所有权
        if doc.owner_id != user_id:
            raise PermissionDeniedError("无权修改此文件")

        # 验证是否为复制文件（复制文件无替换权）
        if doc.is_copied:
            raise PermissionDeniedError("此文件为复制文件，无法替换")

        # 验证文件类型
        if not file.filename:
            raise ValidationError("文件名不能为空")

        file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        if file_ext not in SUPPORTED_FILE_TYPES:
            raise ValidationError(f"不支持的文件类型: {file_ext}")

        # 检查文件大小
        content = await file.read()
        max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
        if len(content) > max_size:
            raise ValidationError(f"文件大小超过限制（最大 {settings.MAX_UPLOAD_SIZE_MB}MB）")

        # 保存新文件
        file_name = f"{uuid.uuid4()}.{file_ext}"
        file_path = os.path.join(settings.UPLOAD_DIR, file_name)

        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        try:
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logger.error(f"文件保存失败: {e}")
            raise ProcessingError("文件保存失败")

        # 保存旧版本到版本历史（验证通过后才保存）
        old_version = DocumentVersion(
            document_id=doc_id,
            version=doc.version,
            filename=doc.filename,
            file_path=doc.file_path,
            file_size=doc.file_size,
            chunk_count=doc.chunk_count,
        )
        db.add(old_version)

        # 删除旧版本的向量数据
        try:
            from app.services.vector_store import get_vector_store
            vector_store = get_vector_store()
            vector_store.delete_documents(doc_id)
        except Exception as e:
            logger.warning(f"删除旧向量数据失败: {e}")

        # 更新文件记录
        doc.filename = file.filename
        doc.file_type = file_ext
        doc.file_size = len(content)
        doc.file_path = file_path
        doc.status = "processing"
        doc.version += 1

        await db.commit()
        await db.refresh(doc)

        # 重新处理文件
        try:
            from app.services.chat import process_document

            chunk_count = await asyncio.to_thread(
                process_document,
                kb_id=doc_id,
                doc_id=doc_id,
                file_path=file_path,
                filename=file.filename,
            )
            doc.status = "completed"
            doc.chunk_count = chunk_count
            logger.info(f"文件替换完成: {file.filename}, 新版本: {doc.version}")
        except Exception as e:
            doc.status = "error"
            error_msg = str(e)
            doc.error_message = error_msg[:500] + "..." if len(error_msg) > 500 else error_msg
            logger.error(f"文件处理失败: {e}")

        await db.commit()
        await db.refresh(doc)

        # 清理旧版本（保留最近3个）
        await DocumentService._cleanup_old_versions(db, doc_id)

        return doc

    @staticmethod
    async def _cleanup_old_versions(db: AsyncSession, doc_id: str):
        """清理旧版本（保留最近3个）"""
        result = await db.execute(
            select(DocumentVersion)
            .where(DocumentVersion.document_id == doc_id)
            .order_by(DocumentVersion.version.desc())
        )
        versions = result.scalars().all()

        # 保留最近3个版本，删除更旧的
        for version in versions[3:]:
            if version.file_path and os.path.exists(version.file_path):
                try:
                    os.remove(version.file_path)
                except Exception:
                    pass
            await db.delete(version)

        await db.commit()

    @staticmethod
    async def generate_description(
        db: AsyncSession, doc_id: str, user_id: str
    ) -> Document:
        """使用 AI 生成文件描述"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 验证所有权
        if doc.owner_id != user_id:
            raise PermissionDeniedError("无权修改此文件")

        # 验证是否为复制文件（复制文件无修改权）
        if doc.is_copied:
            raise PermissionDeniedError("此文件为复制文件，无法修改")

        # 获取文件预览内容
        preview = await DocumentService.get_preview(db, doc_id)
        if not preview:
            raise ValidationError("无法生成描述：文件内容为空")

        # 使用 LLM 生成描述
        try:
            from app.services.llm import get_llm_service

            llm = get_llm_service()
            prompt = f"""请为以下文档生成一个简洁的描述（50-100字），概括文档的主要内容和用途。

文档内容预览：
{preview[:2000]}

请直接输出描述，不要添加任何前缀或解释。"""

            description = await llm.chat(prompt)
            doc.description = description.strip()
            await db.commit()
            await db.refresh(doc)
            logger.info(f"AI 生成文件描述: {doc.filename}")
        except Exception as e:
            logger.error(f"生成文件描述失败: {e}")
            raise ProcessingError("生成描述失败，请稍后重试")

        return doc

    @staticmethod
    async def rollback_version(
        db: AsyncSession, doc_id: str, version: int, user_id: str
    ) -> Document:
        """回滚到指定版本"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 验证所有权
        if doc.owner_id != user_id:
            raise PermissionDeniedError("无权修改此文件")

        # 验证是否为复制文件（复制文件无回滚权）
        if doc.is_copied:
            raise PermissionDeniedError("此文件为复制文件，无法回滚")

        # 查找指定版本
        result = await db.execute(
            select(DocumentVersion).where(
                DocumentVersion.document_id == doc_id,
                DocumentVersion.version == version,
            )
        )
        version_record = result.scalar_one_or_none()
        if not version_record:
            raise NotFoundError("版本", str(version))

        # 删除当前版本的向量数据
        try:
            from app.services.vector_store import get_vector_store
            vector_store = get_vector_store()
            vector_store.delete_documents(doc_id)
        except Exception as e:
            logger.warning(f"删除向量数据失败: {e}")

        # 回滚文件信息
        doc.filename = version_record.filename
        doc.file_path = version_record.file_path
        doc.file_size = version_record.file_size
        doc.status = "processing"

        await db.commit()
        await db.refresh(doc)

        # 重新处理文件
        try:
            from app.services.chat import process_document

            chunk_count = await asyncio.to_thread(
                process_document,
                kb_id=doc_id,
                doc_id=doc_id,
                file_path=doc.file_path,
            )
            doc.status = "completed"
            doc.chunk_count = chunk_count
            logger.info(f"文件回滚完成: {doc.filename}, 版本: {version}")
        except Exception as e:
            doc.status = "error"
            error_msg = str(e)
            doc.error_message = error_msg[:500] + "..." if len(error_msg) > 500 else error_msg
            logger.error(f"文件处理失败: {e}")

        await db.commit()
        await db.refresh(doc)
        return doc

    @staticmethod
    async def add_file(
        db: AsyncSession, doc_id: str, user_id: str
    ) -> dict:
        """添加公开文件到我的文件列表（创建引用关系，不复制文件）"""
        doc = await get_or_404(db, Document, doc_id, "文件")

        # 验证文件是否公开
        if not doc.is_public:
            raise PermissionDeniedError("此文件未公开，无法添加")

        # 不能添加自己的文件
        if doc.owner_id == user_id:
            raise ValidationError("不能添加自己的文件")

        # 检查是否已添加
        existing = await DocumentService._has_user_file(db, user_id, doc_id)
        if existing:
            raise ValidationError("此文件已在您的列表中")

        # 创建用户文件引用
        user_file = UserFile(
            user_id=user_id,
            document_id=doc_id,
        )
        db.add(user_file)

        # 更新原文件的引用次数
        doc.copy_count += 1

        await db.commit()
        logger.info(f"文件已添加: {doc.filename} -> 用户 {user_id}")

        return {"id": doc.id, "filename": doc.filename, "message": "文件已添加到我的文件列表"}

    @staticmethod
    async def get_preview(db: AsyncSession, doc_id: str, max_chars: int = 5000) -> str:
        """获取文件预览内容"""
        doc = await db.get(Document, doc_id)
        if not doc or not doc.file_path:
            return ""

        try:
            if not os.path.exists(doc.file_path):
                return ""

            ext = doc.file_type.lower()

            if ext == "txt":
                with open(doc.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read(max_chars)

            elif ext == "md":
                with open(doc.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read(max_chars)

            elif ext in ("html", "htm"):
                from bs4 import BeautifulSoup
                with open(doc.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    soup = BeautifulSoup(f.read(), "html.parser")
                    return soup.get_text()[:max_chars]

            elif ext == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(doc.file_path)
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    if len("\n".join(text_parts)) > max_chars:
                        break
                return "\n".join(text_parts)[:max_chars]

            elif ext in ("docx", "doc"):
                from docx import Document as DocxDocument
                docx = DocxDocument(doc.file_path)
                paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
                return "\n".join(paragraphs)[:max_chars]

            return ""
        except Exception as e:
            logger.warning(f"获取文件预览失败: {e}")
            return ""
